from io import BytesIO

from docx import Document

from app.services.extractors.base import ExtractionResult
from app.utils.text import clean_text


class DOCXExtractor:
    def extract(self, file_bytes: bytes) -> ExtractionResult:
        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return ExtractionResult(
            text=clean_text("\n".join(paragraphs)),
            detected_format="docx",
            pages=None,
        )
